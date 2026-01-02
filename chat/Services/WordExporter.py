import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_doc_style(doc):
    section = doc.sections[0]
    section.right_to_left = True
    style = doc.styles['Normal']
    style.font.name = 'B Nazanin'
    style.font.size = Pt(12)
    return doc

def summary_to_word(patent_json, output_path):
    doc = Document()
    create_doc_style(doc)


    p = doc.add_paragraph(":  خلاصه اختراع")
    p.runs[0].bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    title = patent_json.get("patent_content", {}).get("invention_title")
    p = doc.add_paragraph(f"عنوان اختراع: {title or 'ارائه نشده'}")
    p.runs[0].bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    abstract = patent_json.get("patent_content", {}).get("abstract", {}).get("text")
    p = doc.add_paragraph(abstract or "ارائه نشده")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(output_path)


def description_to_word(patent_json, output_path):
    doc = Document()
    create_doc_style(doc)

    def add_bold_paragraph(text):
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_to_left = True
        return p

    content = patent_json.get("patent_content", {})
    description = content.get("description", {})
    doc_type = patent_json.get("document_type", "")


    add_bold_paragraph(": توصیف اختراع")
    add_bold_paragraph(": عنوان اختراع به گونه‌ای که در اظهارنامه ذکر گردیده است")


    title = content.get("invention_title", "")
    p = doc.add_paragraph()
    run = p.add_run(title or "ارائه نشده")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_to_left = True


    add_bold_paragraph(": زمینه فنی اختراع مربوط")
    p = doc.add_paragraph()
    run = p.add_run(description.get("technical_field", "ارائه نشده"))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_to_left = True


    add_bold_paragraph(": مشکل فنی و بیان اهداف اختراع")
    p = doc.add_paragraph()
    run = p.add_run(description.get("technical_problem_and_objectives", "ارائه نشده"))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_to_left = True


    add_bold_paragraph(": شرح وضعیت دانش پیشین و سابقه پیشرفت‌های مرتبط")
    p = doc.add_paragraph()
    run = p.add_run(description.get("prior_art_and_background", "ارائه نشده"))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_to_left = True


    add_bold_paragraph(": ارائه راه‌حل و شرح دقیق اختراع")
    text = description.get("solution_and_detailed_description") or \
           description.get("technical_solution", "") + (
               "\n" + description.get("detailed_description", "") if description.get("detailed_description") else "")
    p = doc.add_paragraph()
    run = p.add_run(text or "ارائه نشده")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_to_left = True


    hybrid_expl = description.get("hybrid_structure_explanation")
    if hybrid_expl:
        add_bold_paragraph(": توضیح ساختار ترکیبی")
        p = doc.add_paragraph()
        run = p.add_run(hybrid_expl)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_to_left = True


    process_steps = []
    process_conditions = ""
    materials = ""
    equipment = ""

    if doc_type.startswith("process"):
        process_steps = description.get("process_steps", [])
        process_conditions = description.get("process_conditions", "")
        materials = description.get("materials_and_inputs", "")
        equipment = description.get("equipment_used", "")
    elif doc_type == "hybrid":
        process_specific = content.get("process_specific", {})
        process_steps = process_specific.get("process_stage_sequence", [])
        process_conditions = process_specific.get("operational_conditions", "")
        materials = process_specific.get("required_materials_and_inputs", "")
        equipment = process_specific.get("tools_and_systems_used", "")
        gen_proc = process_specific.get("process_general_description", "")
        if gen_proc:
            add_bold_paragraph(": توصیف کلی فرآیند")
            p = doc.add_paragraph()
            run = p.add_run(gen_proc)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.right_to_left = True

    if process_steps:
        add_bold_paragraph(": شرح مراحل فرآیند تولید")
        for idx, step in enumerate(process_steps, 1):
            if isinstance(step, str):
                p = doc.add_paragraph()
                run = p.add_run(f"{idx}. {step}")
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.right_to_left = True
            elif isinstance(step, dict):
                num = step.get("step_number", idx)
                title = step.get("title", "")
                desc = step.get("description", "")

                p = doc.add_paragraph()
                run = p.add_run(f"{num}. {title}")
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.right_to_left = True

                if desc:
                    p = doc.add_paragraph()
                    run = p.add_run(desc)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.paragraph_format.right_to_left = True
    elif doc_type in ["process", "hybrid"]:
        p = doc.add_paragraph()
        run = p.add_run("ارائه نشده")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_to_left = True

    if any([process_conditions, materials, equipment]):
        add_bold_paragraph(": شرایط فرآیند و مواد/تجهیزات")
        p = doc.add_paragraph()
        run = p.add_run(
            f"شرایط فرآیند: {process_conditions or 'ارائه نشده'}\n"
            f"مواد: {materials or 'ارائه نشده'}\n"
            f"تجهیزات: {equipment or 'ارائه نشده'}"
        )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_to_left = True


    product_specific = content.get("product_specific", {})
    if product_specific:
        gen_desc = product_specific.get("product_general_description", "")
        internal_arch = product_specific.get("product_internal_architecture", "")
        elements = product_specific.get("product_elements_list", [])
        if any([gen_desc, internal_arch, elements]):
            add_bold_paragraph(": مشخصات محصول")
            p = doc.add_paragraph()
            run = p.add_run(f"توصیف کلی محصول: {gen_desc or 'ارائه نشده'}")
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.right_to_left = True

            p = doc.add_paragraph()
            run = p.add_run(f"ساختار داخلی محصول: {internal_arch or 'ارائه نشده'}")
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.right_to_left = True

            if elements:
                add_bold_paragraph(": عناصر و اجزاء محصول")
                for idx, el in enumerate(elements, 1):
                    p = doc.add_paragraph()
                    run = p.add_run(f"{idx}. {el}")
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.paragraph_format.right_to_left = True


    add_bold_paragraph(": مثال اجرایی اختراع")
    p = doc.add_paragraph()
    run = p.add_run(description.get("implementation_examples", "ارائه نشده"))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_to_left = True


    add_bold_paragraph(": مزایای اختراع نسبت به دانش پیشین")
    p = doc.add_paragraph()
    run = p.add_run(description.get("advantages_over_prior_art", "ارائه نشده"))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_to_left = True


    add_bold_paragraph(": کاربرد صنعتی اختراع")
    p = doc.add_paragraph()
    run = p.add_run(description.get("industrial_applicability", "ارائه نشده"))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_to_left = True

    doc.save(output_path)


def claims_to_word(patent_json, output_path):
    doc = Document()
    create_doc_style(doc)

    def add_bold_paragraph(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_to_left = True
        return p

    add_bold_paragraph("ادعانامه")
    add_bold_paragraph(": آنچه ادعا می‌شود")

    claims_data = patent_json.get("patent_content", {}).get("claims", {})
    claims = claims_data.get("independent_claims", []) + claims_data.get("dependent_claims", [])

    if claims:
        for idx, claim in enumerate(claims, 1):
            p = doc.add_paragraph(f"ادعای {idx} {claim}")
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.right_to_left = True
    else:
        p = doc.add_paragraph("ارائه نشده")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_to_left = True

    doc.save(output_path)
